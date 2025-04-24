import os
import requests
import uuid
import json
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Azure Translator settings
subscription_key = '4f2199a9f98a4536892819408c6a7207'  # Your Azure Translator subscription key
endpoint = 'https://api.cognitive.microsofttranslator.com/'  # The endpoint URL for the Translator API
location = 'eastus'  # Your resource location


# API setup
path = '/translate'
constructed_url = endpoint + path


params = {
   'api-version': '3.0',  # API version
   'from': 'en',  # Source language
   'to': ['ar']  # List of target languages
}


headers = {
   'Ocp-Apim-Subscription-Key': subscription_key,
   'Ocp-Apim-Subscription-Region': location,
   'Content-type': 'application/json',
   'X-ClientTraceId': str(uuid.uuid4())
}


def translate_text(text, target_language='ar'):
   body = [{
       'text': text
   }]
   request = requests.post(constructed_url, params=params, headers=headers, json=body)
   response = request.json()
   return response[0]['translations'][0]['text']


def set_paragraph_rtl(paragraph):
   """Set paragraph direction to RTL"""
   p = paragraph._element
   pPr = p.get_or_add_pPr()
   bidi = OxmlElement('w:bidi')
   bidi.set(qn('w:val'), '1')
   pPr.append(bidi)


def translate_document(input_file, output_file, target_language='ar'):
   # Load the document
   doc = Document(input_file)


   # Iterate through paragraphs and translate text
   for paragraph in doc.paragraphs:
       if paragraph.text.strip():  # Skip empty paragraphs
           translated_text = translate_text(paragraph.text, target_language)
           paragraph.text = translated_text
           set_paragraph_rtl(paragraph)  # Set paragraph direction to RTL


   # Iterate through tables and translate text
   for table in doc.tables:
       for row in table.rows:
           for cell in row.cells:
               for paragraph in cell.paragraphs:
                   if paragraph.text.strip():  # Skip empty paragraphs
                       translated_text = translate_text(paragraph.text, target_language)
                       paragraph.text = translated_text
                       set_paragraph_rtl(paragraph)  # Set paragraph direction to RTL


   # Save the translated document
   doc.save(output_file)


def translate_documents_in_folder(input_folder, output_folder, target_language='ar'):
   if not os.path.exists(output_folder):
       os.makedirs(output_folder)


   for filename in os.listdir(input_folder):
       if filename.endswith(".docx"):
           input_file = os.path.join(input_folder, filename)
           output_file = os.path.join(output_folder, filename)
           translate_document(input_file, output_file, target_language)
           print(f"Document {filename} translated and saved to {output_file}")


if __name__ == "__main__":
   input_folder = '/..'      # Input folder path
   output_folder = '/..'     # Output folder path
   target_language = 'ar'  # Arabic


   translate_documents_in_folder(input_folder, output_folder, target_language)
   print(f"All documents in {input_folder} have been translated and saved to {output_folder}")
